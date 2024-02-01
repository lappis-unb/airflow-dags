import pandas as pd
import matplotlib.pyplot as plt
import re
from io import BytesIO
import geopandas as gpd
from docx import Document
import base64
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL

class ReportGenerator:
    def __init__(self):
        pass

    def filter_proposals(self, bp_data, states, pattern):
        proposals = []
        for proposal in bp_data:
            proposal_state = proposal.get("proposal_state")
            proposal_title = proposal.get("proposal_title")
            if proposal_state is not None and proposal_state not in states and bool(re.match(pattern, proposal_title)):
                proposals.append(proposal)
        return proposals

    def create_bp_dataframe(self, bp_data):
        df_bp = pd.DataFrame(bp_data)
        df_bp['publishedAt'] = pd.to_datetime(df_bp['proposal_published_at'])  
        df_bp['updatedAt'] = pd.to_datetime(df_bp['proposal_updated_at']) 
        df_bp['translation'] = df_bp['proposal_title']

        return df_bp

    def calculate_totals(self, bp_data):
        num_proposals = len(bp_data)
        num_votes = sum(proposal['proposal_total_votes'] for proposal in bp_data)
        num_comments = sum(proposal['proposal_total_comments'] for proposal in bp_data)
        return num_proposals, num_votes, num_comments

    def generate_daily_plot(self, df_bp):
        df_bp['proposal_published_at'] = pd.to_datetime(df_bp['proposal_published_at'])
        df_bp['Data'] = df_bp['proposal_published_at'].dt.date

        n_proposals = df_bp.groupby('Data')['proposal_id'].count().to_numpy()
        n_comments = df_bp.groupby('Data')['proposal_total_comments'].sum().to_numpy()
        n_votes = df_bp.groupby('Data')['proposal_total_votes'].sum().to_numpy()

        plt.figure(figsize=(12, 6))
        plt.plot(df_bp['Data'].unique(), n_proposals, label='Propostas', color='blue', marker='o')
        plt.plot(df_bp['Data'].unique(), n_comments, label='Comentários', color='green', marker='s')
        plt.plot(df_bp['Data'].unique(), n_votes, label='Votos', color='red', marker='^')

        plt.xlabel('Data')
        plt.ylabel('Quantidade')
        plt.title('Quantidade de Propostas, Comentários e Votos por Dia')
        plt.legend()
        plt.grid(True)
        plt.xticks(rotation=45)
        plt.tight_layout()

        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)

        daily_graph = base64.b64encode(buffer.getvalue()).decode('utf-8')

        buffer.close()

        return daily_graph
    
    '''def load_data(self, shp_path, json_path):
        brasil = gpd.read_file(shp_path)
        dados = pd.read_json(json_path)
        return brasil, dados

    def filter_and_rename(self, dados, pais, coluna):
        dados_filtrados = dados[dados['country'] == pais]
        dados_filtrados = dados_filtrados.rename(columns={'region': coluna})
        return dados_filtrados

    def create_map(self, brasil, dados, index_coluna, join_coluna):
        mapa = brasil.set_index(index_coluna).join(dados.set_index(join_coluna))
        return mapa

    def plot_map(self, mapa, coluna):
        fig, ax = plt.subplots(figsize=(12, 8))
        mapa.boundary.plot(ax=ax, linewidth=0.5, color='k')
        mapa.plot(column=coluna, ax=ax, legend=True, cmap='YlOrRd')
        plt.title("Visitas por Estado no Brasil")
        plt.axis('off')
        plt.show()

        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)

        map_graph = base64.b64encode(buffer.getvalue()).decode('utf-8')

        buffer.close()

        return map_graph
    
    def devices_plot(json_path, n=3):
        dados = pd.read_json(json_path)
        
        dados = dados.sort_values('nb_visits', ascending=False).head(n)
        
        fig, ax = plt.subplots()
        ax.pie(dados['nb_visits'], labels=dados['label'], autopct='%1.1f%%')
        ax.axis('equal')  
        
        plt.show()

        buffer = BytesIO()
        plt.savefig(buffer, format='png')
        buffer.seek(0)

        dev_graph = base64.b64encode(buffer.getvalue()).decode('utf-8')

        buffer.close()

        return dev_graph'''
    
    def insert_data_docx(self, template_path, data_to_insert, daily_graph):
        doc = Document(template_path)

        table = doc.add_table(rows=1, cols=2)
        for key, value in data_to_insert:
            row_cells = table.add_row().cells
            row_cells[0].text = key
            row_cells[1].text = str(value)

        doc.add_picture(BytesIO(base64.b64decode(daily_graph)), width=Pt(400), height=Pt(300))

        doc.save(template_path)
        